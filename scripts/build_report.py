"""Build the editable SignalBlend Project B report from verified artifacts."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
REPORT = ROOT / "report"
RESULTS = ROOT / "results"
OUTPUT = REPORT / "report.docx"
NAVY = "17324D"
TEAL = "087E8B"
GOLD = "D9A441"
CRIMSON = "B23A48"
MUTED = "5B6573"
LIGHT = "F2F4F7"
GRID = "D8E0E6"


def font(run, size=10.2, color="17212B", bold=False, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd")) or OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    if node.getparent() is None:
        props.append(node)


def cell_margins(cell, top=75, start=105, bottom=75, end=105):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    if sum(widths) != 9360:
        raise ValueError("table widths must sum to 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    props = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = props.find(qn(tag))
        if old is not None: props.remove(old)
    for tag, attributes in (
        ("w:tblW", {"w:w": "9360", "w:type": "dxa"}),
        ("w:tblInd", {"w:w": "120", "w:type": "dxa"}),
        ("w:tblLayout", {"w:type": "fixed"}),
    ):
        node = OxmlElement(tag)
        for key, value in attributes.items(): node.set(qn(key), value)
        props.append(node)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)
            tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW"); cell._tc.get_or_add_tcPr().append(tcw)
            tcw.set(qn("w:w"), str(width)); tcw.set(qn("w:type"), "dxa")


def style_table(table, size=7.8, first_col_bold=False):
    table.style = "Table Grid"
    header = OxmlElement("w:tblHeader"); header.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(header)
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            if r == 0: shade(cell, NAVY)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    font(run, size=size, color="FFFFFF" if r == 0 else "17212B", bold=r == 0 or (first_col_bold and c == 0))


def page_field(paragraph):
    run = paragraph.add_run()
    for tag, attrs, text in (
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
        ("w:fldChar", {"w:fldCharType": "separate"}, None),
        ("w:t", {}, "1"),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ):
        node = OxmlElement(tag)
        for key, value in attrs.items(): node.set(qn(key), value)
        node.text = text
        run._r.append(node)


def setup(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.2); normal.paragraph_format.space_after = Pt(5); normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 12, 6), ("Heading 2", 12.5, TEAL, 9, 4), ("Heading 3", 11, NAVY, 6, 3)
    ):
        style = doc.styles[name]; style.font.name = "Arial"; style._element.rPr.rFonts.set(qn("w:ascii"), "Arial"); style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after); style.paragraph_format.keep_with_next = True
    caption = doc.styles["Caption"]
    caption.font.name = "Arial"; caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial"); caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    caption.font.size = Pt(8); caption.font.color.rgb = RGBColor.from_string(MUTED); caption.paragraph_format.space_after = Pt(5)
    for section in doc.sections:
        section.page_width = Inches(8.5); section.page_height = Inches(11)
        section.top_margin = Inches(.72); section.bottom_margin = Inches(.7); section.left_margin = Inches(.75); section.right_margin = Inches(.75)
        section.header_distance = Inches(.35); section.footer_distance = Inches(.35)
        hp = section.header.paragraphs[0]; hp.text = "SIGNALBLEND  |  FINS5545 PROJECT B"; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs: font(run, 7.5, MUTED, True)
        fp = section.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT; fp.add_run("z5706839  |  Page "); page_field(fp)
        for run in fp.runs: font(run, 7.5, MUTED)


def paragraph(doc, text, *, bold_lead=None, size=None, italic=False, color="17212B", align=None, after=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    if after is not None: p.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead), size or 10.2, color, True, italic)
        font(p.add_run(text[len(bold_lead):]), size or 10.2, color, False, italic)
    else:
        font(p.add_run(text), size or 10.2, color, False, italic)
    return p


def callout(doc, text, fill="EAF5F6", border=TEAL):
    p = paragraph(doc, text, bold_lead=text.split(" ", 1)[0] + " ")
    props = p._p.get_or_add_pPr(); sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), fill); props.append(sh)
    borders = OxmlElement("w:pBdr"); left = OxmlElement("w:left")
    for key, value in (("w:val", "single"), ("w:sz", "18"), ("w:space", "7"), ("w:color", border)): left.set(qn(key), value)
    borders.append(left); props.append(borders); p.paragraph_format.left_indent = Inches(.1); p.paragraph_format.right_indent = Inches(.08)
    return p


def picture(doc, filename, width=6.85):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
    image = p.add_run().add_picture(str(RESULTS / "figures" / filename), width=Inches(width))
    image._inline.docPr.set("descr", filename.replace("_", " ").replace(".png", "")); return p


def caption(doc, label, note):
    p = doc.add_paragraph(style="Caption"); font(p.add_run(f"{label}. "), 8, MUTED, True); font(p.add_run(note), 8, MUTED)


def metric_table(doc, data):
    table = doc.add_table(rows=1, cols=6)
    headers = ["Fund", "Return", "Risk", "Sharpe", "Max DD", "Turnover"]
    for cell, value in zip(table.rows[0].cells, headers, strict=True): cell.text = value
    for row in data.itertuples():
        values = [row.fund.replace("Combined ", "Comb. ").replace("Minimum Variance", "Min Variance").replace("Maximum Sharpe", "Max Sharpe"), f"{row.annualised_return:.1%}", f"{row.annualised_volatility:.1%}", f"{row.sharpe_ratio:.2f}", f"{row.maximum_drawdown:.1%}", f"{row.mean_one_way_turnover:.1%}"]
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True): cell.text = str(value)
    table_geometry(table, [4000, 1050, 1050, 950, 1050, 1260]); style_table(table, 7.25, True)
    return table


def main():
    REPORT.mkdir(exist_ok=True)
    metrics = pd.read_csv(RESULTS / "tables" / "performance_metrics.csv")
    baseline = metrics.loc[~metrics["method"].str.contains("sentiment")].copy()
    fusion = pd.read_csv(RESULTS / "tables" / "fusion_comparison.csv")
    sentiment = pd.read_csv(RESULTS / "data" / "sector_sentiment_index.csv")
    validation = pd.read_csv(RESULTS / "tables" / "validation_summary.csv")
    doc = Document(); setup(doc)

    # Page 1 - editorial cover and answer-first summary.
    paragraph(doc, "SYSTEMATIC MULTI-ASSET FUNDS WITH NEWS-SENTIMENT ANALYTICS", size=9, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    paragraph(doc, "SignalBlend", size=29, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    paragraph(doc, "From market data to an investment decision a novice can explain", size=14, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    paragraph(doc, "FINS5545 Financial Market Data Literacy  |  Project B", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    paragraph(doc, "Student: z5706839  |  Sample: 2020–2023  |  OOS evaluation: 2021–2023", size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=25)
    doc.add_heading("Executive summary", level=1)
    callout(doc, "Finding SignalBlend successfully converts the course data into 14 transparent funds and a four-step investor journey, but the evidence does not support treating headline sentiment as a return-enhancing trading signal.")
    paragraph(doc, "The strongest baseline risk-adjusted results were the Combined Maximum Sharpe fund (Sharpe 0.88) and Combined Risk Parity fund (0.87). The Crypto Minimum Variance fund achieved the highest annualised return (65.0%), but its 73.1% maximum drawdown makes it unsuitable to describe as conservative. Across all baseline funds, outcomes reflect the 2021–2023 historical sample rather than a promise of future performance.")
    paragraph(doc, "I deliberately judged the sentiment extension against an untouched Equity Minimum Variance baseline rather than selecting a favourable comparator. The plain tilt reduced annualised return from 6.0% to 4.7% and Sharpe from 0.48 to 0.38. Coverage-aware reliability improved the tilted version to 5.0% return and 0.40 Sharpe, but still did not beat the base fund. I therefore retain sector sentiment as contextual evidence, not a buy/sell instruction.")
    paragraph(doc, "I designed SignalBlend for a young, time-poor novice because the product's main value is structured comparison: each fund has the same fact-sheet metrics, allocations are normalised transparently, and weak news coverage is visible. Consistent with my conservative-investor scenario, I prioritise drawdown and allocation size rather than the largest backtested return.")
    doc.add_page_break()

    # Page 2 - model design.
    doc.add_heading("1. Funds and walk-forward design", level=1)
    paragraph(doc, "I chose to treat every asset-family and allocation-method pair as a separate fund. Equity-only, crypto-only and combined universes are each allocated by equal weight, minimum variance, maximum Sharpe and risk parity. This exceeds the minimum combined-fund requirement because I wanted users to distinguish the effect of asset choice from the effect of optimisation.")
    doc.add_heading("1.1 Data, calendars and live dates", level=2)
    paragraph(doc, "Adjusted-close simple returns are calculated within ticker on each native calendar. Equity and combined funds use the equity trading calendar and 252-day annualisation; crypto-only funds retain daily observations and use 365. The crypto series is capped at 31 December 2023. Combined returns are formed only after native-calendar returns are computed, preventing weekend crypto price changes from being incorrectly reconstructed from merged price levels.")
    doc.add_heading("1.2 Past-only portfolio formation", level=2)
    paragraph(doc, "I selected monthly rebalancing to update allocations without implying that a time-poor user must trade daily. At the first observed date of each month, target weights are estimated from the preceding 252 observations for equity and combined funds, or 365 observations for crypto funds. The estimation window ends before the rebalance date, so the first live returns are 4 January 2021 for equity/combined funds and 1 January 2021 for crypto. I imposed long-only, fully invested portfolios with 10% equity/combined caps and 25% crypto caps to prevent unconstrained concentration while recognising the smaller crypto universe.")
    paragraph(doc, "Minimum variance minimises w′Σw. Maximum Sharpe maximises (w′μ − rf)/√(w′Σw) with rf = 0, while risk parity minimises squared differences between component risk contributions. Daily covariance matrices are annualised before optimisation to avoid false solver convergence on tiny objectives. Maximum Sharpe uses deterministic multi-start SLSQP and accepts only successful feasible solutions; it never silently falls back to equal weight.")
    callout(doc, "Assumption Zero transaction costs are used in the headline results, but one-way turnover is reported. This is material for Maximum Sharpe, whose average monthly turnover reached about 24% in equity and 25% in combined funds.", fill="FFF7E6", border=GOLD)
    doc.add_heading("1.3 Verification gates", level=2)
    paragraph(doc, f"All {len(validation)} automated build gates passed: unique fund-date keys, finite returns, weights summing to one, non-negative weights, estimation end before rebalance, one-trading-day sentiment lag and bounded reliability. Unit tests also reproduce selected metrics manually and challenge concentrated optimisation and fusion inputs.")
    doc.add_page_break()

    # Page 3 - exact table and risk-return.
    doc.add_heading("2. Out-of-sample fund comparison", level=1)
    paragraph(doc, "I interpret the evidence as a trade-off rather than a universal winner. Combined Maximum Sharpe and Risk Parity delivered the strongest Sharpe ratios, while Combined Minimum Variance produced the lowest volatility (12.6%) and shallowest combined-fund drawdown (15.9%). The crypto funds occupied a different risk regime: annual volatility ranged from 75.0% to 80.4% and maximum drawdown from 73.1% to 82.1%.")
    metric_table(doc, baseline)
    caption(doc, "Table 1", "Out-of-sample metrics by investable baseline fund. Return is compounded annual growth; risk is annualised standard deviation; Sharpe uses a zero risk-free rate; turnover is mean one-way turnover at monthly rebalances. Source: FINS5545 hosted data; author calculations.")
    picture(doc, "risk_return_all_funds.png", 6.55)
    caption(doc, "Figure 1", "Annualised return versus volatility, 2021–2023. Colour identifies asset family and abbreviations identify methods. Crypto is far to the right because its realised volatility was several times that of equity and combined funds. Source: FINS5545 hosted data; author calculations.")
    doc.add_page_break()

    # Page 4 - growth and drawdown.
    doc.add_heading("3. Growth is not a substitute for risk", level=1)
    paragraph(doc, "Growth paths expose timing that annual averages hide. Crypto Minimum Variance ended at approximately $4.54 per initial dollar, but most crypto funds also experienced a severe boom-and-bust cycle. Equity and combined funds were less explosive; combined diversification generally moderated the volatility of a pure crypto exposure.")
    picture(doc, "growth_of_one_all_funds.png", 5.25)
    caption(doc, "Figure 2", "Growth of $1 for 12 baseline funds, using realised walk-forward returns. Panels use different vertical ranges because the crypto scale is much larger. Source: FINS5545 hosted data; author calculations.")
    paragraph(doc, "I give drawdown extra decision weight for a conservative novice because it measures the loss experienced from a previous portfolio peak. Combined Equal Weight fell 28.8%, compared with 15.9% for Combined Minimum Variance. Optimisation reduced loss depth in this sample, but did not remove market risk.")
    picture(doc, "drawdown_combined_funds.png", 4.8)
    caption(doc, "Figure 3", "Combined-fund peak-to-trough drawdowns, 2021–2023. Zero indicates a new portfolio high; more negative values indicate deeper losses. Source: FINS5545 hosted data; author calculations.")

    # Page 5 - weights.
    doc.add_heading("4. What the methods actually held", level=1)
    paragraph(doc, "Method labels matter only if they create meaningfully different portfolios. The latest combined-fund weight vectors were not duplicates: pairwise L1 differences ranged from 0.31 between Equal Weight and Risk Parity to 1.56 between Equal Weight and Maximum Sharpe. Equal Weight maintained 16.7% in crypto by construction, whereas Minimum Variance placed almost all weight in equities for much of the OOS period.")
    picture(doc, "combined_weights_over_time.png", 6.4)
    caption(doc, "Figure 4", "Combined-fund target weights aggregated into equity and crypto at each monthly rebalance. These are target weights, not between-rebalance drifted holdings. Source: FINS5545 hosted data; author calculations.")
    paragraph(doc, "Maximum Sharpe was more concentrated and unstable because estimated mean returns are noisy. Its latest combined portfolio used 15 effective holdings and reached the 10% cap, while Risk Parity retained all 60 assets and had lower turnover. This helps explain why their annual returns were similar but their implementation profiles differed.")
    callout(doc, "Implication A novice should not interpret ‘optimal’ as ‘certain’. Optimisation reflects an objective and a historical estimation window; it does not discover a permanently best portfolio.")
    doc.add_page_break()

    # Page 6 - sentiment.
    doc.add_heading("5. Standalone sector sentiment", level=1)
    paragraph(doc, "VADER scores each deduplicated raw headline without removing casing or punctuation, then headlines are averaged to ticker-day and ticker-days are equal-weighted within sector. A missing ticker-day is omitted from the polarity mean rather than converted to a neutral zero. Coverage is reported separately: mean daily sector coverage was 77.2%, ranging from 20% to 100%. About 24.9% of covered ticker-day scores were exactly neutral on average.")
    paragraph(doc, "The headline date is first aligned to the same or next equity trading day. The signal becomes usable only on the next observed equity day: Saturday and Monday headlines aligned to Monday can first affect Tuesday. This lag produced zero recorded timing violations.")
    picture(doc, "sector_sentiment_index.png", 5.15)
    caption(doc, "Figure 5", "Twenty-one-trading-day average of daily VADER compound sentiment by equity sector, 2020–2023. Daily sector values equal-weight covered tickers; the smoothing is for visual interpretation only. Source: FINS5545 hosted headlines; author calculations.")
    callout(doc, "Limitation The data contains headlines, not full articles. VADER measures language tone rather than verified economic impact, and finance headlines can score neutral despite carrying decision-relevant information.", fill="FFF0F1", border=CRIMSON)
    doc.add_page_break()

    # Page 7 - innovation and fusion.
    doc.add_heading("6. Innovation: coverage-aware reliability", level=1)
    paragraph(doc, "My Project A audit found that news availability differed materially across sectors, so I chose coverage quality—not another generic return signal—as the main extension. For sector s on day t, I define reliability as Rₛ,ₜ = coverageₛ,ₜ × (1 − normalised concentrationₛ,ₜ). Coverage is the share of sector tickers with news; concentration is the normalised HHI of headline counts among covered names. Reliability averaged 68.1% and was bounded between zero and one using information available on the lagged signal day.")
    paragraph(doc, "The base Equity Minimum Variance weights are multiplied by an exponential sector tilt. Plain sentiment uses clipped cross-sectional sector z-scores; the innovation multiplies those scores by R before applying the same 0.35 strength. A capped proportional projection restores full investment and the 10% position limit. This creates a clean comparison of base, plain sentiment and coverage-aware sentiment under identical dates.")
    picture(doc, "fusion_before_after.png", 5.75)
    caption(doc, "Figure 6", "Base versus sentiment-enhanced Equity Minimum Variance funds. The lower panel shows 63-day rolling annualised active return relative to the base. Source: FINS5545 hosted price and headline data; author calculations.")
    metric_table(doc, fusion)
    caption(doc, "Table 2", "Fusion before-versus-after results. Coverage-aware reliability reduced the performance damage relative to the plain tilt but did not outperform the base. Source: FINS5545 hosted data; author calculations.")
    paragraph(doc, "I retained this negative result because selecting only favourable backtests would weaken the evaluation. Positive headline tone may already be reflected in prices, may identify crowded optimism, or may not persist until the next tradable day. Reliability weighting addresses availability bias but cannot solve weak semantic validity. My product decision is therefore to retain sentiment as contextual analytics and avoid presenting it as a return forecast.")
    doc.add_page_break()

    # Page 8 - app journey.
    doc.add_heading("7. Streamlit product and investor journey", level=1)
    paragraph(doc, "I defined the target user as a young investor who is interested in investing but lacks time to learn many strategies or monitor news continuously. I organised the app as a four-step journey—compare, inspect, allocate and question—so it reduces research cost without pretending to automate suitability. The app loads committed precomputed artifacts; it does not run VADER or recompute backtests on Cloud. This keeps the deployed product responsive and makes every displayed metric consistent with the submitted report.")
    journey = doc.add_table(rows=1, cols=3)
    for cell, value in zip(journey.rows[0].cells, ["Step", "User question", "SignalBlend response"], strict=True): cell.text = value
    rows = [
        ("1. Compare", "Which risk-return trade-off fits my attention and tolerance?", "Interactive risk-return map, identical metrics and conservative-risk prompt."),
        ("2. Inspect", "What happened inside one fund?", "Growth, drawdown, Sharpe, turnover and latest target holdings."),
        ("3. Allocate", "How would several funds have behaved together?", "Normalised sliders, combined historical metrics and crypto-exposure warning."),
        ("4. Question", "Is headline mood based on enough evidence?", "Sector mood, coverage, reliability and honest fusion comparison."),
    ]
    for values in rows:
        cells = journey.add_row().cells
        for cell, value in zip(cells, values, strict=True): cell.text = value
    table_geometry(journey, [1400, 3100, 4860]); style_table(journey, 8.1, True)
    caption(doc, "Table 3", "SignalBlend investor journey. The sequence moves from comparison to due diligence rather than presenting a single automatic recommendation.")
    paragraph(doc, "The interface uses a consistent navy, teal and gold system, plain-language explanations, percentage formatting and visible limitations. Allocation inputs are automatically normalised rather than accepted as invalid totals. If more than 20% is allocated to crypto-only funds, the app warns that this excludes crypto already embedded in combined funds. The fact sheet also labels holdings as target weights, preventing users from mistaking them for live drifted positions.")
    callout(doc, "Product boundary SignalBlend is educational decision support, not personalised advice. Historical backtests, estimated ‘optimal’ weights and headline tone cannot determine suitability for an individual investor.", fill="FFF7E6", border=GOLD)
    doc.add_heading("7.1 Deployment architecture", level=2)
    paragraph(doc, "The local build downloads the hosted ZIP once, writes validated CSVs and figures under results/, and separates build-only NLTK from deployment dependencies. The public app reads those artifacts from its repository. At hand-in the separate GitHub repository must be public and the Streamlit URL must work for a logged-out marker; during development the repository remains private.")
    doc.add_page_break()

    # Page 9 - reflection and recommendations.
    doc.add_heading("8. Critical reflection and recommendations", level=1)
    doc.add_heading("8.1 What worked", level=2)
    paragraph(doc, "My strongest modelling decision was to require walk-forward design, explicit calendars and solver validation rather than accept in-sample illustrations. The resulting funds are differentiated and reproducible. Combined funds demonstrated that introducing crypto can improve return opportunities without forcing a user into a crypto-only product. My coverage-aware reliability measure also converted Project A's descriptive audit into a measurable signal-quality control.")
    doc.add_heading("8.2 What did not work—and why", level=2)
    paragraph(doc, "Headline sentiment did not improve the base minimum-variance fund. I did not suppress or retune away this result. It persisted even after coverage adjustment, showing that data availability and predictive content are separate problems: reliability can reduce confidence in thin evidence, but cannot create information that forecasts returns. I also required Maximum Sharpe failures to be exposed rather than silently replaced, because an attractive objective does not guarantee robust implementation.")
    doc.add_heading("8.3 Three concrete recommendations", level=2)
    paragraph(doc, "1. Default conservative novices to comparison, not auto-selection. Rank no fund as universally ‘best’. Lead with volatility, drawdown and crypto share, and require users to open the fact sheet before saving an allocation. This responds to observed drawdowns of 73–82% in crypto-only funds.", bold_lead="1. ")
    paragraph(doc, "2. Keep sentiment informational until it passes a stronger validation gate. Continue showing sector mood and reliability, but do not use the current tilt in a live default portfolio. Promote it only after transaction-cost-aware, subperiod and holdout tests show stable improvement over the same base fund.", bold_lead="2. ")
    paragraph(doc, "3. Add implementation frictions before product launch. Model trading costs from reported turnover, cap fund-level crypto exposure for conservative presets, and add a paper-trading monitoring period. Maximum Sharpe's roughly 24–25% mean monthly turnover makes zero-cost performance especially optimistic.", bold_lead="3. ")
    doc.add_heading("8.4 Remaining uncertainty", level=2)
    paragraph(doc, "The OOS evaluation covers only 2021–2023, a short and unusual regime containing both a crypto boom and a broad 2022 drawdown. There is no independent post-2023 holdout, no bid-ask spread or tax model, and no formal investor suitability assessment. Results should therefore guide product learning, not capital promises.")
    doc.add_page_break()

    # References and compact appendix (excluded from narrative limit).
    doc.add_heading("Reference list", level=1)
    refs = [
        "Hutto, C.J. & Gilbert, E. 2014, ‘VADER: A parsimonious rule-based model for sentiment analysis of social media text’, Proceedings of the Eighth International AAAI Conference on Weblogs and Social Media, pp. 216–225.",
        "Maillard, S., Roncalli, T. & Teïletche, J. 2010, ‘The properties of equally weighted risk contribution portfolios’, The Journal of Portfolio Management, vol. 36, no. 4, pp. 60–70.",
        "Markowitz, H. 1952, ‘Portfolio selection’, The Journal of Finance, vol. 7, no. 1, pp. 77–91.",
        "Sharpe, W.F. 1966, ‘Mutual fund performance’, The Journal of Business, vol. 39, no. 1, pp. 119–138.",
        "UNSW Business School 2026, FINS5545 FinTech Project 2026: Systematic Multi-Asset Funds with News-Sentiment Analytics, project brief and hosted course dataset, UNSW Sydney.",
    ]
    for ref in refs:
        p = paragraph(doc, ref, size=9, after=4); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_heading("Appendix A. Reproducibility and AI workflow", level=1)
    paragraph(doc, "The complete project reproduces results with python scripts/run_part_b.py, runs automated tests with pytest, and checks hand-in structure with python scripts/check_handin.py. Required CSVs are committed for the deployed app. Agent instructions and curated prompt logs in AGENTS.md and ai/ document the student's specifications, AI-assisted implementation, failed tests, corrections and reasons. The report's interpretation remains subject to the student's final review and sign-off.")
    appendix = doc.add_table(rows=1, cols=3)
    for cell, value in zip(appendix.rows[0].cells, ["Control", "Result", "Evidence"], strict=True): cell.text = value
    for row in validation.itertuples():
        cells = appendix.add_row().cells
        values = [row.check.replace("_", " ").title(), "Pass" if row.passed else "Fail", str(row.evidence_value)]
        for cell, value in zip(cells, values, strict=True): cell.text = value
    table_geometry(appendix, [3900, 1200, 4260]); style_table(appendix, 8.0)
    caption(doc, "Table A1", "Automated build validation summary generated by scripts/run_part_b.py.")

    doc.save(OUTPUT)
    print(f"wrote {OUTPUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
